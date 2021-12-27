import base64
import datetime
import io
import os
from datetime import date, datetime
from functools import reduce

import docx
import numpy as np
import pandas as pd

import otchet_oitscb.load_data as ld
import otchet_oitscb.params as params
from otchet_oitscb import log_writer as lw


def bound_replace(string, old):
    return string.replace(old, '')


def parse_contents(contents, filename):
    content_type, content_string = contents.split(',')
    decoded = base64.b64decode(content_string)
    try:
        content_df = pd.read_excel(io.BytesIO(decoded), skiprows=7).drop(0)
        content_df = ld.load_df(df=content_df)
        lw.log_writer(log_msg=f'Файл {filename} успешно загружен')

        return content_df
    except Exception as e:
        lw.log_writer(log_msg=f'При загрузки файла "{filename}" возникла ошибка: {e}')
        content_df = ld.no_data()

        return content_df


def count_params(df, start_date, end_date, area, tech_group, total=True):
    if total:
        df = df[(df.reg_date >= datetime(date.today().year, 1, 1)) & (df.reg_date <= end_date)]
        df = df[df.assign_group == params.oitscb_group].reset_index()
        df.drop('index', inplace=True, axis=1)
        total_count = df[df.func_area.isin(area)]['task_number'].count()
        not_solve = df[df.func_area.isin(area)].groupby('status')['task_number'].count()[['В ожидании', 'В работе',
                                                                                          'Открыт']].sum()
        inf_inq = df[df.func_area.isin(area)].groupby(['status', 'wait'])['task_number'].count()[
            'В ожидании', 'Запрос информации у пользователя'].sum()

        return total_count, not_solve, inf_inq

    else:
        df.reg_date = pd.to_datetime(df.reg_date)
        df = df[(df.reg_date >= start_date) & (df.reg_date <= end_date)]

        total_count = df[(df.func_area.isin(area)) & (df.assign_group == params.oitscb_group)]['task_number'].count()
        not_solve = df[(df.func_area.isin(area)) & (df.assign_group == params.oitscb_group)]. \
            groupby('status')['task_number'].count()[['В ожидании', 'В работе', 'Открыт']].sum()
        inf_inq = df[(df.func_area.isin(area)) & (df.assign_group == params.oitscb_group)].groupby(['status', 'wait'])[
            'task_number'].count()['В ожидании', 'Запрос информации у пользователя'].sum()
        in_work = not_solve - inf_inq
        work_3l = df[(df.func_area.isin(area)) & (df.assign_group == params.oitscb_group)].groupby(['status', 'wait'])[
            'task_number'].count()['В ожидании', 'Работа в рамках дочернего запроса'].sum()
        work_2l = in_work - work_3l
        work_tech = df[(df.func_area.isin(area)) & (df.assign_group.isin([params.tech_group, tech_group]))]['task_number'].count()

        return total_count, not_solve, inf_inq, in_work, work_2l, work_3l, work_tech


def build_table(curr_data, prev_data, total_data, curr_period, prev_period):
    res_df2 = pd.DataFrame(columns=['  ', f'{prev_period}', f'{curr_period}', 'Изменение за неделю'])
    res_df2.loc[0] = 'кол-во обращений', prev_data[0], curr_data[0], (
            (curr_data[0] - prev_data[0]) / prev_data[0]) * 100
    res_df2.loc[1] = 'всего с начала 2021г.', np.nan, total_data[0], np.nan
    res_df2.loc[2] = 'не решено', prev_data[1], curr_data[1], ((curr_data[1] - prev_data[1]) / prev_data[1]) * 100
    res_df2.loc[3] = 'не решено, %', round(prev_data[1] / prev_data[0] * 100, 2), round(
        curr_data[1] / curr_data[0] * 100, 2), np.nan
    res_df2.loc[4] = 'всего с начала 2021г', np.nan, total_data[1], np.nan
    res_df2.loc[5] = 'на запросе информации у пользователя', prev_data[2], curr_data[2], np.nan
    res_df2.loc[6] = 'на запросе информации у пользователя, %', round(prev_data[2] / prev_data[1] * 100, 2), round(
        curr_data[2] / curr_data[1] * 100, 2), np.nan
    res_df2.loc[7] = 'в работе', prev_data[3], curr_data[3], np.nan
    res_df2.loc[8] = 'из них на 2-й линии', prev_data[4], curr_data[4], np.nan
    res_df2.loc[9] = 'из них на 3-й линии', prev_data[5], curr_data[5], np.nan
    res_df2.loc[10] = 'в работе у группы технологов', prev_data[6], curr_data[6], np.nan
    res_df2['Изменение за неделю'] = res_df2['Изменение за неделю'].round(2)
    res_df2 = res_df2.fillna('-')
    return res_df2


def count_params_categories(df, start_date, end_date, category_list, name_list, sub):
    df = df[(df.reg_date >= start_date) & (df.reg_date <= end_date)]
    df = df[df.assign_group == params.oitscb_group].reset_index()
    df.drop('index', inplace=True, axis=1)

    df.analytics = df.analytics.fillna('-')
    df.analytics = df.analytics.apply(
        lambda x: reduce(bound_replace, ['Сложность>;', '>', '<', 'Сложность: ', 'Категория обращения: '], x))
    df.analytics = df.analytics.apply(lambda x: x.split(';'))

    df['category'] = df.analytics.apply(lambda row: [item for item in row if not item.strip().isdigit()])
    df['category'] = df['category'].apply(lambda x: [x[i].strip() for i in range(len(x))])
    df['category'] = df['category'].apply(lambda x: x[0] if len(x) > 0 else x)
    df['category'] = df['category'].apply(lambda x: '' if x == [] else x)
    df['category'] = df['category'].apply(lambda x: 'Не указано' if x == '' else x)

    df['main_category'] = ''
    for i in range(len(category_list)):
        df.loc[df[df['category'].isin(category_list[i])].index, 'main_category'] = name_list[i]

    df['subsystem'] = ''
    for i in range(len(category_list)):
        df.loc[df[df['category'].isin(category_list[i])].index, 'subsystem'] = sub

    cat_df = df.pivot_table(index=['subsystem', 'main_category'], values='task_number', aggfunc='count').reset_index()

    cat_df = cat_df.merge(df.pivot_table(index=['subsystem'],
                                         values=['task_number'],
                                         aggfunc={'task_number': 'count'}
                                         ).reset_index().rename(columns={'task_number': 'total_tasks'}))

    cat_df['persent'] = round(cat_df['task_number'] / cat_df['total_tasks'] * 100, 2)

    cat_df.drop(['subsystem', 'total_tasks'], axis=1, inplace=True)
    cat_df.drop(0, axis=0, inplace=True)

    return cat_df


def build_table_categories(df, prev_df, curr_period, prev_period):
    df.rename(columns={'task_number': f'{curr_period}', 'persent': 'Доля в общем количестве, %'}, inplace=True)
    prev_df.rename(columns={'task_number': f'{prev_period}', 'persent': 'Доля в общем количестве, %.'},
                   inplace=True)

    merge_df = prev_df.merge(df, on=['main_category'], how='outer')
    merge_df.rename(columns={'main_category': 'Категория обращения'}, inplace=True)
    merge_df = merge_df.fillna('-')

    return merge_df


def write_to_word(header_list, df_list):
    doc = docx.Document()
    doc.add_heading('Статистика по сопровождению', 0)

    doc.add_heading(header_list[0], 1)
    table = doc.add_table(rows=len(df_list[0])+1, cols=4)
    table.style = 'Table Grid'
    for column_id, column_name in enumerate(df_list[0].columns):
        cell = table.cell(0, column_id)
        cell.text = column_name

    for row in range(1, len(df_list[0])+1):
        for col in range(4):
            cell = table.cell(row, col)
            cell.text = str(df_list[0].iloc[row-1, col])
    doc.add_page_break()

    doc.add_heading(header_list[1], 1)
    table = doc.add_table(rows=len(df_list[1]) + 1, cols=len(df_list[1].columns))
    table.style = 'Table Grid'
    for column_id, column_name in enumerate(df_list[1].columns):
        cell = table.cell(0, column_id)
        cell.text = column_name

    for row in range(1, len(df_list[1]) + 1):
        for col in range(len(df_list[1].columns)):
            cell = table.cell(row, col)
            cell.text = str(df_list[1].iloc[row - 1, col])
    doc.add_page_break()

    doc.add_heading(header_list[2], 1)
    table = doc.add_table(rows=len(df_list[2]) + 1, cols=4)
    table.style = 'Table Grid'
    for column_id, column_name in enumerate(df_list[2].columns):
        cell = table.cell(0, column_id)
        cell.text = column_name

    for row in range(1, len(df_list[2]) + 1):
        for col in range(4):
            cell = table.cell(row, col)
            cell.text = str(df_list[2].iloc[row - 1, col])
    doc.add_page_break()

    doc.add_heading(header_list[3], 1)
    table = doc.add_table(rows=len(df_list[3]) + 1, cols=len(df_list[3].columns))
    table.style = 'Table Grid'
    for column_id, column_name in enumerate(df_list[3].columns):
        cell = table.cell(0, column_id)
        cell.text = column_name

    for row in range(1, len(df_list[3]) + 1):
        for col in range(len(df_list[3].columns)):
            cell = table.cell(row, col)
            cell.text = str(df_list[3].iloc[row - 1, col])
    doc.add_page_break()

    if os.path.exists(params.filepath):
        doc.save(os.path.join(params.filepath, params.filename))

    else:
        lw.log_writer('Отсутствует каталог для файлов с отчетами. Создаем каталог...')
        os.mkdir(params.filepath)
        lw.log_writer('Каталог создан успешно')
        doc.save(os.path.join(params.filepath, params.filename))


def create_report(df, prev_df, start_date, end_date, prev_date, area, tech_group):

    curr_period = ' '.join([start_date.strftime('%d-%m-%Y'), end_date.strftime('%d-%m-%Y')])
    prev_period = ' '.join([prev_date.strftime('%d-%m-%Y'), start_date.strftime('%d-%m-%Y')])

    curr_data = count_params(
        df=df,
        start_date=start_date,
        end_date=end_date,
        area=area,
        tech_group=tech_group,
        total=False)
    prev_data = count_params(
        df=prev_df,
        start_date=prev_date,
        end_date=start_date,
        area=area,
        tech_group=tech_group,
        total=False)
    total_data = count_params(
        df=df,
        start_date=start_date,
        end_date=end_date,
        area=area,
        tech_group=tech_group,
        total=True)
    data_df = build_table(
        curr_data=curr_data,
        prev_data=prev_data,
        total_data=total_data,
        curr_period=curr_period,
        prev_period=prev_period
    )
    return data_df


def create_report_cat(df, prev_df, start_date, end_date, prev_date, cat_list, name_list, sub):

    curr_period = ' '.join([start_date.strftime('%d-%m-%Y'), end_date.strftime('%d-%m-%Y')])
    prev_period = ' '.join([prev_date.strftime('%d-%m-%Y'), start_date.strftime('%d-%m-%Y')])

    curr_data = count_params_categories(
        df=df,
        start_date=start_date,
        end_date=end_date,
        category_list=cat_list,
        name_list=name_list,
        sub=sub
    )
    prev_data = count_params_categories(
        df=prev_df,
        start_date=prev_date,
        end_date=start_date,
        category_list=cat_list,
        name_list=name_list,
        sub=sub
    )

    data_df = build_table_categories(
        df=curr_data,
        prev_df=prev_data,
        curr_period=curr_period,
        prev_period=prev_period
        )
    return data_df
